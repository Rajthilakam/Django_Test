from __future__ import unicode_literals
from django.shortcuts import render,redirect,render_to_response
from django.template import RequestContext
import docx
from docx import Document
import re
import io
import xml.etree.ElementTree as ET
from collections import OrderedDict
from dicttoxml import dicttoxml
from xml.etree.ElementTree import Element,SubElement,Comment,tostring
from .forms import CandidateForm
from .import models
from .models import Candidate
from django.contrib import messages
from django.conf import settings
from django.views.static import serve

def index(request):
    return render(request,'uploadresume.html')


def results(request):
    if request.method == 'POST':
        filedoc = request.FILES['filename']
        wordDoc = Document(request.FILES['filename'])
        dict1 = OrderedDict()
        valuelist = []
        mylist = ['firstname', 'lastname', 'contactdetails', 'emailid', 'addressdetails', 'workexperience',
                  'visastatus', 'workpermit', 'readytorelocate', 'technology']
        for paragraphs in wordDoc.paragraphs:
            line = paragraphs.text
            valuelist = line.split(':')
            if len(valuelist) > 1:
                dict1[valuelist[0].strip()] = valuelist[1].strip()
            check_list = list(dict1.keys())
        l = [item.lower() for item in check_list]
        if l == mylist:
            for paragraphs in wordDoc.paragraphs:
                line = paragraphs.text
                valuelist = line.split(':')
                if len(valuelist) > 1:
                    dict1[valuelist[0].strip()] = valuelist[1].strip()
            finalxml = dicttoxml(dict1, attr_type=False)
            root = ET.fromstring(finalxml)


# Create your views here.
#Assigning values to the models
            candidateDetail = Candidate()
            candidateDetail.firstname   = root[0].text
            candidateDetail.lastname    = root[1].text
            candidateDetail.contactdetails = root[2].text
            candidateDetail.emailid   = root[3].text
            candidateDetail.addressdetails = root[4].text
            candidateDetail.workexperience = root[5].text
            candidateDetail.visastatus = root[6].text
            candidateDetail.workpermit = root[7].text
            candidateDetail.readytorelocate = root[8].text
            candidateDetail.technology = root[9].text
            candidateDetail.save()
    #context = {'candidateDetail': candidateDetail,}
            print ("candidate class")
            print (candidateDetail.lastname)
            return render(request,'candidatedetail.html',{"candidateDetail":candidateDetail})

        else:
            messages.error(request,"Invalid Resume format.Your resume should be in the following format.click the link below")
            return render(request,'uploadresume.html')

def  update_candidatedetail(request, id):
    candidateDetail = Candidate.objects.get(id=id)
    form = CandidateForm(request.POST, instance=candidateDetail)
    if request.method == 'POST':
        print("success")
        if form.is_valid():
            form.save()
            print ("Form updated")
    return render(request,'candidatedetail.html',{"candidateDetail":candidateDetail})
